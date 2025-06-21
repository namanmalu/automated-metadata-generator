import os
import re
import docx
import fitz  # PyMuPDF
import pytesseract
from pdf2image import convert_from_path
from collections import Counter
import spacy
from en_core_web_sm import load as load_model
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity

nlp = load_model()

SECTION_HEADINGS = [
    "abstract", "introduction", "objective", "problem",
    "conclusion", "summary", "results", "discussion"
]

def get_named_entities(text):
    doc = nlp(text)
    entities = {}
    people = list(set([ent.text for ent in doc.ents if ent.label_ == "PERSON"]))
    orgs = list(set([ent.text for ent in doc.ents if ent.label_ == "ORG"]))
    gpes = list(set([ent.text for ent in doc.ents if ent.label_ == "GPE"]))
    dates = list(set([ent.text for ent in doc.ents if ent.label_ == "DATE"]))

    if people:
        entities["people"] = people
    if orgs:
        entities["organizations"] = orgs
    if gpes:
        entities["locations"] = gpes
    if dates:
        entities["dates"] = dates

    return entities

def extract_sections(text):
    lines = text.split("\n")
    current_section = None
    sections = {}
    for line in lines:
        clean_line = line.strip()
        line_lower = clean_line.lower()
        if any(heading in line_lower for heading in SECTION_HEADINGS):
            clean_heading = re.sub(r"^[0-9]+[.)]?\s*", "", clean_line)
            current_section = clean_heading
            sections[current_section] = ""
        elif current_section:
            sections[current_section] += clean_line + " "
    return sections

def get_top_sentences(text, n=5):
    sentences = re.split(r'(?<=[.!?]) +', text)
    if len(sentences) <= n:
        return sentences
    tfidf = TfidfVectorizer(stop_words='english')
    tfidf_matrix = tfidf.fit_transform(sentences)
    sim_scores = cosine_similarity(tfidf_matrix[0:1], tfidf_matrix).flatten()
    top_indices = sim_scores.argsort()[-n:][::-1]
    return [sentences[i] for i in top_indices if i < len(sentences)]

def structure_metadata(metadata_dict):
    structured = {}
    for key, value in metadata_dict.items():
        if isinstance(value, list) and value:
            structured[key] = [str(v).strip() for v in value if str(v).strip()]
        elif isinstance(value, dict) and value:
            structured[key] = {k: v for k, v in value.items() if v}
        elif value not in [None, "", "None"]:
            structured[key] = str(value).strip()
    return structured

def extract_smart_metadata(text, filename=""):
    words = re.findall(r'\b\w{4,}\b', text.lower())
    common_words = Counter(words).most_common(10)
    top_sentences = get_top_sentences(text)
    entities = get_named_entities(text)
    sections = extract_sections(text)

    summary = sections.get("summary", "") or sections.get("conclusion", "") or " ".join(top_sentences[:2])
    objective = sections.get("objective", "") or sections.get("introduction", "") or " ".join(top_sentences[2:4])

    metadata = {
        "filename": filename,
        "word_count": len(words),
        "character_count": len(text),
        "top_keywords": [word for word, _ in common_words],
        "key_sentences": top_sentences,
        "summary": summary.strip(),
        "purpose": objective.strip(),
        "named_entities": entities if entities else None,
        "sections_found": list(sections.keys())
    }

    return structure_metadata(metadata)

def extract_docx_metadata(docx_path):
    doc = docx.Document(docx_path)
    full_text = "\n".join([para.text for para in doc.paragraphs])
    title = doc.paragraphs[0].text.strip() if doc.paragraphs else "Untitled"
    metadata = extract_smart_metadata(full_text, os.path.basename(docx_path))
    metadata["title"] = title
    return structure_metadata(metadata)

def extract_pdf_metadata(pdf_path):
    text = ""
    try:
        with fitz.open(pdf_path) as doc:
            for page in doc:
                page_text = page.get_text()
                if page_text.strip():
                    text += page_text
        if not text.strip():
            images = convert_from_path(pdf_path)
            for img in images:
                text += pytesseract.image_to_string(img)
    except Exception as e:
        text = f"Error reading PDF: {str(e)}"
    return structure_metadata(extract_smart_metadata(text, os.path.basename(pdf_path)))

def extract_txt_metadata(txt_path):
    with open(txt_path, 'r', encoding='utf-8') as f:
        text = f.read()
    return structure_metadata(extract_smart_metadata(text, os.path.basename(txt_path)))

def extract_metadata(file_path):
    ext = os.path.splitext(file_path)[1].lower()
    if ext == ".pdf":
        return extract_pdf_metadata(file_path)
    elif ext == ".docx":
        return extract_docx_metadata(file_path)
    elif ext == ".txt":
        return extract_txt_metadata(file_path)
    else:
        return structure_metadata({
            "error": "Unsupported file type",
            "filename": os.path.basename(file_path)
        })

def convert_metadata_to_csv(metadata):
    from io import StringIO
    import csv

    output = StringIO()
    writer = csv.writer(output)
    writer.writerow(["Field", "Value"])
    for key, value in metadata.items():
        if isinstance(value, list):
            value = ", ".join(map(str, value))
        elif isinstance(value, dict):
            value = ", ".join([f"{k}: {', '.join(v)}" for k, v in value.items()])
        writer.writerow([key, value])
    return output.getvalue()


